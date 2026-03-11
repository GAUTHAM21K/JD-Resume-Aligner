import sys
import re
import os
import requests
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor

# UI Enhancements
import typer
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table

load_dotenv()

app = typer.Typer(help="🚀 Resume Tailor 2026")
console = Console()

GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')
OUTPUT_DIR = Path("tailored_resumes")
OUTPUT_DIR.mkdir(exist_ok=True)

def read_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading {file_path}: {e}"

def save_as_docx(text_content, output_filename):
    doc = Document()
    # Remove common AI markdown artifacts that break clean formatting
    clean_content = text_content.replace("**", "").replace("__", "").replace("###", "").strip()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    for line in clean_content.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        
        # Detect Headings: If line is short and all caps, or specific keywords
        headers = ["PROFILE SUMMARY", "CORE COMPETENCIES", "SKILLS", "PROJECTS", "EXPERIENCE", "ACADEMICS", "EDUCATION"]
        if line.isupper() or any(h in line.upper() for h in headers) and len(line) < 30:
            h = doc.add_paragraph()
            run = h.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith(('*', '-', '•')):
            p = doc.add_paragraph(line.lstrip('*-•').strip(), style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)

    save_path = OUTPUT_DIR / output_filename
    doc.save(save_path)
    return save_path

def get_ai_response(prompt: str):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3-flash-preview:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.2, "maxOutputTokens": 8000} # Increased tokens
    }
    
    try:
        response = requests.post(url, json=payload, timeout=90)
        if response.status_code != 200:
            url_alt = f"https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"
            response = requests.post(url_alt, json=payload, timeout=90)
        
        response.raise_for_status()
        return response.json()['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        return None

@app.command()
def start(
    resume_path: str = typer.Option("master/resume.docx", "--resume", "-r"),
    projects_dir: str = typer.Option("projects/", "--projects", "-p")
):
    console.print(Panel("RESUME TAILOR 2026", subtitle="Deep Optimization Mode", style="bold cyan"))
    
    console.print("[bold yellow]📋 Paste Job Description:[/bold yellow] [dim](Ctrl+D/Z when done)[/dim]")
    jd_text = sys.stdin.read().strip()

    if not jd_text:
        console.print("[red]❌ JD empty.[/red]")
        raise typer.Exit()

    with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}"), transient=True) as progress:
        progress.add_task(description="Processing files...", total=None)
        resume_text = read_docx(resume_path)
        
        projects_context = ""
        p_path = Path(projects_dir)
        if p_path.exists():
            for p in p_path.glob("*"):
                if p.suffix in ['.txt', '.md', '.docx']:
                    content = read_docx(str(p)) if p.suffix == '.docx' else p.read_text(errors='ignore')
                    projects_context += f"\n--- PROJECT FILE: {p.name} ---\n{content}\n"

        progress.add_task(description="AI Tailoring (STAR Method)...", total=None)
        prompt = f"""
        You are an expert Resume Writer. Your goal is to produce a FULL, COMPLETE resume tailored to the JD.

        CONTEXT:
        - Job Description: {jd_text}
        - Master Resume: {resume_text}
        - Available Projects: {projects_context}

        CRITICAL CONSTRAINTS:
        1. OUTPUT THE FULL RESUME from start to finish. Do not truncate.
        2. KEEP EXACTLY: Contact Info, EXPERIENCE section, and ACADEMICS section from the Master Resume.
        3. TAILOR: 'PROFILE SUMMARY', 'SKILLS', and 'PROJECTS'.
        4. STAR METHOD: Rewrite project bullet points as: [Action Verb] [Task] using [Tool] resulting in [Quantifiable Result].
        5. NO MARKDOWN: Do not use **bold** or __underline__ syntax. Use plain text only.
        
        STRUCTURE:
        COMPANY_NAME: [Identify Company]
        ### START_RESUME ###
        [Full Name & Contact]
        PROFILE SUMMARY
        [Tailored Summary]
        CORE COMPETENCIES
        [Tailored Skills List]
        EXPERIENCE
        [Master Experience - Unchanged]
        PROJECTS
        [Tailored Projects from Project Repository/Resume]
        ACADEMICS
        [Master Academics - Unchanged]
        ### END_RESUME ###
        MATCH_SCORE: [0-100]
        """
        
        full_response = get_ai_response(prompt)
        if not full_response:
            raise typer.Exit(1)

        # Better Parsing Logic
        try:
            company_name = re.search(r"COMPANY_NAME:\s*(.*)", full_response).group(1).strip()
            content = re.search(r"### START_RESUME ###(.*)### END_RESUME ###", full_response, re.DOTALL).group(1).strip()
            score = re.search(r"MATCH_SCORE:\s*(.*)", full_response).group(1).strip()
        except:
            content = full_response
            company_name = "Tailored_Resume"
            score = "Check Output"

        output_file = f"Resume_{company_name.replace(' ', '_')}.docx"
        final_path = save_as_docx(content, output_file)

    console.print(f"\n[bold green]✅ Done! Saved to {final_path}[/bold green]")
    console.print(f"[bold cyan]Match Score: {score}[/bold cyan]")

if __name__ == "__main__":
    app()